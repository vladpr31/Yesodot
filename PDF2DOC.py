from pdf2image import convert_from_bytes
from PyPDF2 import PdfFileWriter, PdfFileReader
import io
import sys
from PIL import Image
from docx import Document

if len(sys.argv) > 1:
    input_file = sys.argv[1]
else:
    input_file = input("\x1b[0;33mType the PDF file name to create cover photo:\x1b[0m ")

inp = PdfFileReader(input_file, "rb")
page = inp.getPage(0)

wrt = PdfFileWriter()
wrt.addPage(page)

r = io.BytesIO()
wrt.write(r)

images = convert_from_bytes(r.getvalue())
images[0].save(input_file[:-4] + ".jpg")
print("\x1b[0;33mDone, your cover photo has been saved as {}.\x1b[0m".format(input_file[:-4] + ".jpg"))
r.close()


def crop(image_path, coords, saved_location):
    """
    @param image_path: The path to the image to edit
    @param coords: A tuple of x/y coordinates (x1, y1, x2, y2)
    @param saved_location: Path to save the cropped image
    """
    image_obj = Image.open(image_path)
    cropped_image = image_obj.crop(coords)
    cropped_image.save(saved_location)
    document = Document()
    tables = document.tables
    table = document.add_table(rows=1, cols=2)
    row_cells = table.add_row().cells

    for i, image in enumerate(['cropped3.jpg']):
        paragraph = row_cells[i].paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(image)
    document.save('cropped3.docx')
    x = input("Would you like to open the Picture?(y/n) ")
    if x == 'y':
        cropped_image.show()
    else:
        return None


if __name__ == '__main__':
    image = '2.jpg'
    crop(image, (321, 204, 1543, 546), 'cropped3.jpg')
    """first 2 variables indicate top left coords and 2 last indicate
    and the last 2 indicate the right bottom coords both (x,y)&(x,y)"""
